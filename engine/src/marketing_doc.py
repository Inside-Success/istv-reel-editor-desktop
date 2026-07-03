"""Render client-facing marketing package (HTML + Word)."""

from __future__ import annotations

import html
import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ISTV_SHOWS = frozenset({
    "istv-people",
    "istv-legacymakers",
    "istv-operationceo",
    "istv-womeninpower",
    "istv-mompreneurs",
})

TABLE_HEADERS = ("#", "Reel Title", "Caption + Hashtags", "SEO Title", "Post When", "Type")

# Inside Success / ISTV palette (from reference marketing package)
COLOR_INK = RGBColor(34, 26, 43)
COLOR_GOLD = RGBColor(230, 180, 80)
COLOR_SLATE = RGBColor(36, 26, 48)
COLOR_MUTED = RGBColor(107, 100, 117)


def _text(value) -> str:
    return str(value if value is not None else "").strip()


def _reel_label(reel: dict) -> str:
    rank = int(reel.get("rank") or reel.get("id") or 0)
    title = _text(reel.get("title"))
    if not title:
        return str(rank)
    return f"{rank}. {title}"


def _caption_cell(reel: dict) -> str:
    caption = _text(reel.get("caption"))
    tags = " ".join(_text(h) for h in (reel.get("hashtags") or []) if _text(h))
    if tags and not tags.startswith("#"):
        tags = " ".join(f"#{t.lstrip('#')}" for t in tags.split())
    if caption and tags:
        return f"{caption}\n\n{tags}"
    return caption or tags


def _hashtags_html(reel: dict) -> str:
    tags = [_text(h) for h in (reel.get("hashtags") or []) if _text(h)]
    if not tags:
        return ""
    spans = []
    for tag in tags:
        clean = tag if tag.startswith("#") else f"#{tag.lstrip('#')}"
        spans.append(f"<span>{html.escape(clean)}</span>")
    return f'<div class="tags">{"".join(spans)}</div>'


def resolve_istv_collaboration(recommendations: dict) -> str:
    if not isinstance(recommendations, dict):
        return ""
    direct = _text(recommendations.get("istv_collaboration"))
    if direct:
        return _normalize_show_handle(direct)
    single = recommendations.get("istv_show")
    if isinstance(single, dict):
        return _normalize_show_handle(single.get("account"))
    if isinstance(single, str):
        return _normalize_show_handle(single)
    niches = recommendations.get("niche_accounts") or []
    if isinstance(niches, list) and niches:
        first = niches[0] if isinstance(niches[0], dict) else {}
        return _normalize_show_handle(first.get("account"))
    return ""


def _normalize_show_handle(value: str) -> str:
    handle = _text(value).lower()
    if not handle:
        return ""
    handle = handle.replace("@", "")
    if not handle.startswith("istv"):
        return ""
    if handle in ISTV_SHOWS:
        return handle
    for show in ISTV_SHOWS:
        if show in handle:
            return show
    return handle


def _doc_stem(result: dict, doc_title: str) -> str:
    title = _text(doc_title)
    if "—" in title:
        return title.split("—", 1)[0].strip()
    return title or "Documentary"


def render_marketing_doc_html(
    result: dict,
    doc_title: str = "Documentary — Short-Form Marketing Package",
) -> str:
    reels = sorted(result.get("reels") or [], key=lambda r: int(r.get("rank") or r.get("id") or 0))
    summary = html.escape(_text(result.get("documentary_summary")))
    show = resolve_istv_collaboration(result.get("recommendations") or {}) or "istv-people"
    stem = html.escape(_doc_stem(result, doc_title))
    series_parts = [r for r in reels if r.get("series_part")]
    series_note = ""
    if series_parts:
        ids = ", ".join(str(int(r.get("rank") or r.get("id") or 0)) for r in series_parts)
        series_note = (
            f'<div class="series-note"><span class="pip"></span>'
            f"<span><b>{len(series_parts)}-part series</b> across Reels {ids} — post in order.</span></div>"
        )

    reel_cards = []
    for reel in reels:
        rank = int(reel.get("rank") or reel.get("id") or 0)
        rank_label = f"{rank:02d}"
        title = html.escape(_text(reel.get("title")))
        caption = html.escape(_text(reel.get("caption")))
        seo = html.escape(_text(reel.get("seo_title")))
        content_type = html.escape(_text(reel.get("content_type")) or "Reel moment")
        is_series = bool(reel.get("series_part"))
        series_class = " series" if is_series else ""
        ribbon = (
            f'<span class="ribbon">Series · Part {int(reel["series_part"])}</span>'
            if is_series
            else ""
        )
        posting = html.escape(_text(reel.get("best_posting_time")) or "TBD")
        score = int(reel.get("score") or 0)
        score_badge = f'<span class="score">{score}</span>' if score else ""
        reel_cards.append(
            f"""
      <article class="reel{series_class}">
        <div class="slate">
          <div class="slate-top"><span class="lbl">Reel</span><span class="no">{rank_label}</span>{score_badge}</div>
          <div class="slate-tag">{content_type}</div>
          <div class="post-when"><span class="l">Post</span>{posting}</div>
        </div>
        <div class="body">
          <div class="reel-head">
            <h3 class="reel-title">{title}</h3>
            {ribbon}
          </div>
          <p class="caption">{caption}</p>
          <div class="seo"><span class="l">SEO</span><em>{seo}</em></div>
          {_hashtags_html(reel)}
        </div>
      </article>"""
        )

    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{stem} — Short-Form Marketing Package</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Archivo:wght@600;700;800&family=Inter:wght@400;500;600&family=IBM+Plex+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>
  :root{{
    --ink:#1A1224;--slate:#1E1428;--paper:#FFFFFF;--panel:#F0ECF5;--line:#DDD6E8;
    --gold:#E6B450;--mint:#79E6C1;--muted:#6B6475;--muted-2:#9A92A6;--series:#6D5BD0;--maxw:1100px;
  }}
  *{{box-sizing:border-box}}
  body{{margin:0;background:linear-gradient(180deg,#EDE8F4 0%,#F7F5FA 240px,#F7F5FA 100%);color:var(--ink);font-family:Inter,system-ui,sans-serif;line-height:1.55}}
  .wrap{{max-width:var(--maxw);margin:0 auto;padding:0 24px}}
  .hero{{background:linear-gradient(135deg,#1E1428 0%,#2A1A38 55%,#1A2830 100%);color:#F4EFF8;position:relative;overflow:hidden;border-bottom:4px solid var(--gold)}}
  .hero::before{{content:"";position:absolute;inset:0;background:radial-gradient(ellipse 80% 60% at 85% 15%,rgba(230,180,80,.18),transparent 55%)}}
  .hero .wrap{{padding:56px 24px 48px;position:relative;z-index:1}}
  .brand-mark{{font-family:"IBM Plex Mono",monospace;font-size:11px;letter-spacing:.28em;text-transform:uppercase;color:var(--gold);margin:0 0 20px;display:flex;align-items:center;gap:10px}}
  .brand-mark::before{{content:"";width:28px;height:2px;background:var(--mint)}}
  h1{{font-family:Archivo,sans-serif;font-weight:800;font-size:clamp(32px,5.5vw,48px);line-height:1.08;margin:0 0 14px;max-width:18ch}}
  .lede{{max-width:58ch;font-size:16px;color:#C8BED8;margin:0;line-height:1.65}}
  .meta{{display:flex;flex-wrap:wrap;gap:10px;margin-top:28px}}
  .chip{{font-family:"IBM Plex Mono",monospace;font-size:11px;border:1px solid rgba(255,255,255,.2);color:#E8E1F2;border-radius:8px;padding:7px 13px;background:rgba(255,255,255,.04)}}
  .chip b{{color:var(--gold);font-weight:600}}
  .recs{{margin:-28px auto 0;position:relative;z-index:2}}
  .recs .wrap{{display:grid;grid-template-columns:1fr;gap:16px}}
  .card-pad{{background:var(--paper);border:1px solid var(--line);border-radius:18px;padding:24px 26px;box-shadow:0 20px 50px -28px rgba(26,18,36,.35)}}
  .rec-label{{font-family:"IBM Plex Mono",monospace;font-size:10px;letter-spacing:.2em;text-transform:uppercase;color:var(--muted);margin:0 0 12px}}
  .collab{{display:flex;align-items:center;gap:16px}}
  .collab .num{{width:48px;height:48px;border-radius:12px;background:var(--slate);color:var(--gold);font-family:Archivo;font-weight:800;font-size:20px;display:flex;align-items:center;justify-content:center;flex-shrink:0}}
  .collab h3{{font-family:Archivo;font-weight:700;font-size:18px;margin:0;color:var(--ink)}}
  .collab p{{margin:5px 0 0;font-size:13.5px;color:var(--muted)}}
  .series-note{{margin-top:16px;font-size:13px;color:var(--muted);display:flex;gap:10px;align-items:flex-start;padding:12px 14px;background:#F5F2FA;border-radius:10px}}
  .series-note .pip{{width:8px;height:8px;border-radius:50%;background:var(--series);margin-top:6px;flex-shrink:0}}
  .sec{{display:flex;align-items:baseline;justify-content:space-between;margin:48px 0 20px}}
  .sec h2{{font-family:Archivo;font-weight:700;font-size:22px;margin:0;color:var(--ink)}}
  .sec span{{font-family:"IBM Plex Mono",monospace;font-size:11px;color:var(--muted-2);letter-spacing:.06em}}
  .reels{{display:grid;grid-template-columns:repeat(auto-fill,minmax(480px,1fr));gap:16px;padding-bottom:48px}}
  .reel{{background:var(--paper);border:1px solid var(--line);border-radius:16px;display:grid;grid-template-columns:118px 1fr;overflow:hidden;box-shadow:0 8px 24px -16px rgba(26,18,36,.12);transition:box-shadow .2s}}
  .reel:hover{{box-shadow:0 16px 36px -18px rgba(26,18,36,.2)}}
  .slate{{background:var(--slate);color:#EDE7F5;padding:16px 12px;display:flex;flex-direction:column;gap:12px;position:relative}}
  .slate::before{{content:"";position:absolute;top:0;left:0;right:0;height:3px;background:var(--mint)}}
  .reel.series .slate::before{{background:var(--series)}}
  .slate-top{{display:flex;align-items:baseline;flex-wrap:wrap;gap:4px}}
  .slate-top .lbl{{font-family:"IBM Plex Mono",monospace;font-size:9px;letter-spacing:.18em;text-transform:uppercase;color:var(--muted-2);width:100%}}
  .slate-top .no{{font-family:Archivo;font-weight:800;font-size:32px;color:var(--gold);line-height:1}}
  .reel.series .slate-top .no{{color:#C9BEF2}}
  .score{{font-family:"IBM Plex Mono",monospace;font-size:10px;background:rgba(121,230,193,.15);color:var(--mint);border-radius:6px;padding:2px 6px;margin-left:auto}}
  .slate-tag{{font-family:"IBM Plex Mono",monospace;font-size:9.5px;color:#C9C0DA;border:1px solid rgba(255,255,255,.14);border-radius:6px;padding:5px 6px;text-align:center;line-height:1.3}}
  .post-when{{font-size:10px;color:#B8AED0;line-height:1.35}}
  .post-when .l{{font-family:"IBM Plex Mono",monospace;font-size:8px;letter-spacing:.14em;text-transform:uppercase;display:block;color:var(--muted-2);margin-bottom:2px}}
  .body{{padding:18px 20px 20px}}
  .reel-head{{display:flex;align-items:flex-start;gap:10px;justify-content:space-between}}
  .reel-title{{font-family:Archivo;font-weight:700;font-size:17px;line-height:1.25;margin:0;color:var(--ink)}}
  .ribbon{{font-family:"IBM Plex Mono",monospace;font-size:9px;text-transform:uppercase;color:#5a49b8;background:#EEEAFB;border:1px solid #DCD4F6;border-radius:999px;padding:4px 9px;white-space:nowrap;flex-shrink:0}}
  .caption{{margin:10px 0 0;font-size:14px;color:#3d3548;white-space:pre-line}}
  .seo{{margin:12px 0 0;font-size:12.5px;color:var(--muted);display:flex;gap:8px;align-items:flex-start;padding:10px 12px;background:#FAF8FC;border-radius:8px;border-left:3px solid var(--gold)}}
  .seo .l{{font-family:"IBM Plex Mono",monospace;font-size:9px;letter-spacing:.12em;text-transform:uppercase;color:var(--muted-2);flex-shrink:0;padding-top:2px}}
  .seo em{{font-style:normal;color:var(--ink)}}
  .tags{{margin-top:12px;display:flex;flex-wrap:wrap;gap:5px}}
  .tags span{{font-size:11px;color:#1e5c45;background:#E4F7EF;border:1px solid #C5EBD9;border-radius:6px;padding:3px 8px;font-family:"IBM Plex Mono",monospace}}
  footer{{margin:8px 0 56px;text-align:center;padding-top:24px;border-top:1px solid var(--line)}}
  footer p{{font-family:"IBM Plex Mono",monospace;font-size:11px;color:var(--muted-2);margin:0}}
  @media (max-width:560px){{.reels{{grid-template-columns:1fr}}.reel{{grid-template-columns:88px 1fr}}}}
</style>
</head>
<body>
  <header class="hero">
    <div class="wrap">
      <p class="brand-mark">Inside Success · ISTV</p>
      <h1>{stem}</h1>
      <p class="lede">{summary}</p>
      <div class="meta">
        <span class="chip"><b>{len(reels)}</b> reels</span>
        <span class="chip">recommended <b>@{html.escape(show)}</b></span>
      </div>
    </div>
  </header>
  <section class="recs"><div class="wrap">
    <div class="card-pad">
      <p class="rec-label">Best ISTV Reel Collaboration</p>
      <div class="collab"><div class="num">★</div><div><h3>@{html.escape(show)}</h3><p>Primary niche account for this documentary package.</p></div></div>
      {series_note}
    </div>
  </div></section>
  <main class="wrap">
    <div class="sec"><h2>Posting Package</h2><span>{len(reels)} reels · captions · SEO · schedule</span></div>
    <div class="reels">{"".join(reel_cards)}</div>
  </main>
  <footer class="wrap"><p>Inside Success · ISTV Short-Form Marketing Package</p></footer>
</body>
</html>"""


def _set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def _add_heading(doc: Document, text: str, *, size: int = 16, color: RGBColor = COLOR_INK) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    para.paragraph_format.space_after = Pt(10)


def _add_body(doc: Document, text: str, *, bold: bool = False, space_after: int = 8) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_INK
    para.paragraph_format.space_after = Pt(space_after)


def _shade_header_row(row) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_SLATE


def render_marketing_doc_docx(
    result: dict,
    output_path: str | Path,
    doc_title: str = "Documentary — Short-Form Marketing Package",
) -> str:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(render_marketing_doc_bytes(result, doc_title=doc_title))
    html_path = path.with_suffix(".html")
    html_path.write_text(render_marketing_doc_html(result, doc_title=doc_title), encoding="utf-8")
    return str(path.resolve())


def _add_meta_line(doc: Document, label: str, value: str) -> None:
    para = doc.add_paragraph()
    label_run = para.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = COLOR_MUTED
    value_run = para.add_run(value)
    value_run.font.size = Pt(11)
    value_run.font.color.rgb = COLOR_INK
    para.paragraph_format.space_after = Pt(6)


def render_marketing_doc_bytes(
    result: dict,
    doc_title: str = "Documentary — Short-Form Marketing Package",
) -> bytes:
    buffer = io.BytesIO()
    doc = Document()
    _set_document_defaults(doc)

    # Title block
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(doc_title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = COLOR_SLATE
    title_para.paragraph_format.space_after = Pt(4)

    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run("ISTV Short-Form Marketing Package")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = COLOR_GOLD
    sub_run.bold = True
    sub_para.paragraph_format.space_after = Pt(14)

    summary = _text(result.get("documentary_summary"))
    if summary:
        _add_body(doc, summary, space_after=16)

    show = resolve_istv_collaboration(result.get("recommendations") or {}) or "istv-people"
    _add_heading(doc, "Collaboration & Distribution", size=13, color=COLOR_SLATE)
    _add_meta_line(doc, "Recommended ISTV account", f"@{show}")
    reels = sorted(result.get("reels") or [], key=lambda r: int(r.get("rank") or r.get("id") or 0))
    series_parts = [r for r in reels if r.get("series_part")]
    if series_parts:
        ids = ", ".join(str(int(r.get("rank") or r.get("id") or 0)) for r in series_parts)
        _add_meta_line(doc, "Series", f"{len(series_parts)}-part sequence across Reels {ids}")

    doc.add_paragraph()
    _add_heading(doc, f"Reel Package ({len(reels)} posts)", size=13, color=COLOR_SLATE)

    table = doc.add_table(rows=1, cols=len(TABLE_HEADERS))
    table.style = "Table Grid"
    table.autofit = False
    widths = (0.45, 1.35, 2.4, 1.5, 0.95, 0.85)
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    for idx, label in enumerate(TABLE_HEADERS):
        cell = table.rows[0].cells[idx]
        cell.text = label
    _shade_header_row(table.rows[0])

    for reel in reels:
        rank = int(reel.get("rank") or reel.get("id") or 0)
        row = table.add_row().cells
        row[0].text = f"{rank:02d}"
        row[1].text = _text(reel.get("title"))
        row[2].text = _caption_cell(reel)
        row[3].text = _text(reel.get("seo_title"))
        row[4].text = _text(reel.get("best_posting_time")) or "TBD"
        ctype = _text(reel.get("content_type")) or "reel"
        if reel.get("series_part"):
            ctype = f"{ctype} · Part {reel['series_part']}"
        row[5].text = ctype

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = footer.add_run("Inside Success · ISTV Short-Form Marketing Package")
    foot_run.font.size = Pt(9)
    foot_run.font.color.rgb = COLOR_MUTED
    foot_run.italic = True

    doc.save(buffer)
    return buffer.getvalue()


def normalize_recommendations(raw: dict | None) -> dict:
    raw = raw if isinstance(raw, dict) else {}
    show = resolve_istv_collaboration(raw) or "istv-people"
    return {"istv_collaboration": show}
